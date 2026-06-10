from __future__ import annotations

import json
import mimetypes
from pathlib import Path

from bs4 import BeautifulSoup
from docx import Document
from pypdf import PdfReader


###############################################################################
def resolve_local_path(path_value: str) -> Path:
    return Path(path_value).expanduser().resolve()


###############################################################################
def _html_to_text(raw_html: str) -> str:
    soup = BeautifulSoup(raw_html, "html.parser")
    for tag in soup(["script", "style", "nav", "header", "footer", "aside"]):
        tag.decompose()
    return "\n".join(
        part.strip() for part in soup.get_text("\n").splitlines() if part.strip()
    )


###############################################################################
def _load_docx_text(path: Path) -> str:
    document = Document(str(path))
    return "\n".join(
        paragraph.text.strip()
        for paragraph in document.paragraphs
        if paragraph.text.strip()
    )


###############################################################################
def load_docx_paragraphs(path: Path) -> list[dict[str, object]]:
    document = Document(str(path))
    return [
        {"paragraph_index": index + 1, "text": paragraph.text.strip()}
        for index, paragraph in enumerate(document.paragraphs)
        if paragraph.text.strip()
    ]


###############################################################################
def _load_pdf_text(path: Path) -> str:
    reader = PdfReader(str(path))
    parts: list[str] = []
    for page in reader.pages:
        extracted = (page.extract_text() or "").strip()
        if extracted:
            parts.append(extracted)
    return "\n\n".join(parts)


###############################################################################
def load_pdf_pages(
    path: Path, *, include_empty_pages: bool = False
) -> list[dict[str, object]]:
    reader = PdfReader(str(path))
    pages: list[dict[str, object]] = []
    for index, page in enumerate(reader.pages, start=1):
        extracted = (page.extract_text() or "").strip()
        if extracted or include_empty_pages:
            pages.append({"page_number": index, "text": extracted})
    return pages


###############################################################################
def _read_text_file(path: Path) -> str:
    try:
        return path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        return path.read_text(encoding="latin-1")


###############################################################################
def load_file_text(path: Path) -> tuple[str, str]:
    suffix = path.suffix.lower()
    if suffix in {".txt", ".md", ".markdown"}:
        return _read_text_file(path), mimetypes.guess_type(str(path))[0] or "text/plain"
    if suffix in {".html", ".htm"}:
        raw = _read_text_file(path)
        return _html_to_text(raw), mimetypes.guess_type(str(path))[0] or "text/html"
    if suffix == ".json":
        payload = json.loads(_read_text_file(path))
        text_payload = (
            payload
            if isinstance(payload, str)
            else json.dumps(payload, indent=2, ensure_ascii=True, default=str)
        )
        return text_payload, "application/json"
    if suffix in {".csv", ".tsv", ".log", ".xml", ".yaml", ".yml"}:
        return _read_text_file(path), mimetypes.guess_type(str(path))[0] or "text/plain"
    if suffix == ".rtf":
        return _read_text_file(path), "application/rtf"
    if suffix == ".pdf":
        return _load_pdf_text(path), "application/pdf"
    if suffix == ".docx":
        return _load_docx_text(
            path
        ), "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    raise ValueError(f"Unsupported document type: {path.suffix or path.name}")


__all__ = [
    "load_docx_paragraphs",
    "load_file_text",
    "load_pdf_pages",
    "resolve_local_path",
]
